from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/venele/Downloads/fault-injection")
INPUT = ROOT / "output/doc/云平台故障注入工具-模板填充版.docx"
OUTPUT = ROOT / "output/doc/云平台故障注入工具-终稿格式精修版.docx"
BACKUP = ROOT / "output/doc/云平台故障注入工具-模板填充版-终稿精修前备份.docx"


RECENT_REFERENCES = [
    "[1] Al-Said Ahmad A, Al-Qora'n L F, Zayed A. Exploring the impact of chaos engineering with various user loads on cloud native applications: an exploratory empirical study[J]. Computing, 2024, 106: 2389-2425.",
    "[2] Al-Said Ahmad A, Andras P. Scalability resilience framework using application-level fault injection for cloud-based software services[J]. Journal of Cloud Computing, 2022, 11(1): 1.",
    "[3] Mailewa A B, Akuthota A, Mohottalalage T M D. A review of resilience testing in microservices architectures: implementing chaos engineering for fault tolerance and system reliability[C]//2025 IEEE 15th Annual Computing and Communication Workshop and Conference. 2025: 236-242.",
    "[4] Yadav R. Harnessing chaos: the role of chaos engineering in cloud applications and impacts on site reliability engineering[J]. International Journal of Computer Trends and Technology, 2024, 72(6): 25-30.",
    "[5] Poltronieri F, Tortonesi M, Stefanelli C. ChaosTwin: a chaos engineering and digital twin approach for the design of resilient IT services[C]//2021 17th International Conference on Network and Service Management. 2021: 234-238.",
    "[6] Naqvi M A, Malik S, Astekin M, Moonen L. On evaluating self-adaptive and self-healing systems using chaos engineering[C]//2022 IEEE International Conference on Autonomic Computing and Self-Organizing Systems. 2022: 1-10.",
    "[7] Torkura K A, Sukmana M I H, Cheng F, Meinel C. CloudStrike: chaos engineering for security and resiliency in cloud infrastructure[J]. IEEE Access, 2020, 8: 123044-123060.",
    "[8] Cloud Native Computing Foundation. Chaos engineering in 2024 with LitmusChaos[EB/OL]. 2024.",
    "[9] Chaos Mesh Authors. Chaos Mesh overview[EB/OL]. 2026.",
    "[10] Chaos Mesh Authors. Basic features[EB/OL]. 2026.",
    "[11] Microsoft. Faults and actions in Azure Chaos Studio[EB/OL]. 2024.",
    "[12] Microsoft. Azure Chaos Studio fault and action library[EB/OL]. 2026.",
    "[13] The Linux Kernel Documentation. Kernel probes (Kprobes)[EB/OL]. 2024.",
    "[14] The Linux Kernel Documentation. Fault injection capabilities infrastructure[EB/OL]. 2026.",
    "[15] The Linux Kernel Documentation. The definitive KVM API documentation[EB/OL]. 2026.",
    "[16] The Linux Kernel Documentation. ARM KVM documentation[EB/OL]. 2026.",
    "[17] QEMU Project. QEMU version 9.2.0 released[EB/OL]. 2024.",
    "[18] Apache CloudStack. Apache CloudStack 4.20.0.0 release[EB/OL]. 2024.",
    "[19] Apache CloudStack. CloudStack API documentation (4.20.0.0)[EB/OL]. 2024.",
    "[20] Apache Hadoop. Apache Hadoop 3.4.1[EB/OL]. 2024.",
    "[21] Kubernetes Authors. Horizontal pod autoscaling[EB/OL]. 2026.",
    "[22] Gremlin. Fault injection[EB/OL]. 2026.",
]


CITATION_RULES = [
    ("随着云计算技术的持续发展", "[1-3,21]"),
    ("故障注入技术通过人为构造可控异常", "[1-2,11-12,22]"),
    ("故障注入作为系统可靠性与容错性评测的重要方法", "[1-7]"),
    ("在经典研究中，FIAT", "[1-8,22]"),
    ("结合当前云平台运行环境的新特点", "[13-20]"),
    ("在典型的云计算架构中", "[15-17,21]"),
    ("平台管理栈负责实例创建", "[18-20]"),
    ("故障注入是指在可控条件下", "[1-2,9-14]"),
    ("从故障表现形式上看", "[9-14,22]"),
    ("云平台中的容错行为通常具有分布式和动态性的特点", "[1-3,21]"),
    ("为了更贴近当前云平台运行环境的实际特点", "[15-17]"),
    ("在虚拟化层技术选型方面", "[15-17]"),
    ("在关键技术选型方面", "[13-20]"),
    ("平台设计的总体目标可以概括为四个方面", "[1-6,9-12]"),
    ("故障模型管理模块负责把平台支持的不同故障类型", "[9-14]"),
    ("任务描述不仅包含故障模型中的核心字段", "[1-3,11-12]"),
    ("很多故障测试工具只关注如何把故障注入进去", "[5-6,9-12]"),
    ("Kprobes 是 Linux 提供的一种轻量级动态探测框架", "[13]"),
    ("在 KVM 虚拟化环境中", "[15-17]"),
    ("内存管理故障是云平台底层最具代表性的高风险异常之一", "[13-16]"),
    ("除计算、内存与存储路径外", "[15-17]"),
    ("CPU 资源争抢与内存耗尽注入", "[11-14,22]"),
    ("网络异常是分布式系统最常见", "[9-12,18-20,22]"),
    ("Hadoop 集群由 NameNode", "[20]"),
    ("CloudStack 作为典型云管理平台", "[18-19]"),
    ("后端控制中枢是平台控制面的核心部分", "[9-12,18-19]"),
    ("实验指标的设计需要同时覆盖故障注入工具本身", "[1-3]"),
    ("通过上述实验，可以从多个角度验证平台的有效性", "[1-6]"),
    ("未来可在现有任务调度基础上增加实验模板库", "[5-6,9-12]"),
]


PAGE_START_TITLES = {
    "摘  要",
    "Abstract",
    "目  录",
    "结  论",
    "参考文献",
    "攻读学士学位期间取得创新性成果",
    "原创性声明和使用权限",
    "致  谢",
}


POST_REFERENCE_TITLES = {
    "攻读学士学位期间取得创新性成果",
    "原创性声明和使用权限",
    "致  谢",
}


def set_run_font(run, east_asia: str = "宋体", latin: str = "Times New Roman", size: float | None = None, bold: bool | None = None) -> None:
    run.font.name = latin
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)
    run._element.rPr.rFonts.set(qn("w:ascii"), latin)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold


def format_para_runs(para: Paragraph, east_asia: str = "宋体", latin: str = "Times New Roman", size: float = 12, bold: bool | None = None) -> None:
    if not para.runs:
        para.add_run("")
    for run in para.runs:
        set_run_font(run, east_asia=east_asia, latin=latin, size=size, bold=bold)


def set_line_spacing(para: Paragraph, multiple: float) -> None:
    """Write explicit WordprocessingML spacing so Word does not drop caption/reference line rules."""
    spacing = para._p.get_or_add_pPr().find(qn("w:spacing"))
    if spacing is None:
        spacing = OxmlElement("w:spacing")
        para._p.get_or_add_pPr().append(spacing)
    spacing.set(qn("w:line"), str(int(round(240 * multiple))))
    spacing.set(qn("w:lineRule"), "auto")


def ensure_page_break_before(para: Paragraph) -> None:
    para.paragraph_format.page_break_before = True


def add_keep_next(para: Paragraph) -> None:
    ppr = para._p.get_or_add_pPr()
    if ppr.find(qn("w:keepNext")) is None:
        ppr.append(OxmlElement("w:keepNext"))


def remove_keep_next(para: Paragraph) -> None:
    ppr = para._p.get_or_add_pPr()
    keep = ppr.find(qn("w:keepNext"))
    if keep is not None:
        ppr.remove(keep)


def citation_insert(text: str, cite: str) -> str:
    if cite in text:
        return text
    stripped = text.rstrip()
    if stripped.endswith("。"):
        return stripped[:-1] + cite + "。"
    return stripped + cite


def remove_paragraph(para: Paragraph) -> None:
    parent = para._element.getparent()
    if parent is not None:
        parent.remove(para._element)


def insert_paragraph_after(doc: Document, current_el, text: str, style: str = "Plain Text") -> Paragraph:
    new_p = OxmlElement("w:p")
    current_el.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    para.style = style
    para.text = text
    return para


def is_main_chapter(text: str) -> bool:
    return bool(re.match(r"^第[1-9]\d*章\s+", text))


def is_figure_caption(text: str) -> bool:
    return bool(re.match(r"^图\d+-\d+\s+", text))


def is_table_caption(text: str) -> bool:
    return bool(re.match(r"^表\d+-\d+\s+", text))


def set_no_table_borders(table) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = qn(f"w:{edge}")
        element = borders.find(tag)
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "nil")


def clear_paragraph_indent(para: Paragraph) -> None:
    para.paragraph_format.first_line_indent = None
    para.paragraph_format.left_indent = None
    para.paragraph_format.right_indent = None


def format_front_matter(doc: Document) -> None:
    """Keep cover pages out of the generic body formatter."""
    try:
        abstract_idx = next(i for i, para in enumerate(doc.paragraphs) if para.text.strip() == "摘  要")
    except StopIteration:
        return

    nonempty = [(i, para) for i, para in enumerate(doc.paragraphs[:abstract_idx]) if para.text.strip()]
    if len(nonempty) < 9:
        return

    cover_type = nonempty[0][1]
    cover_cn_title = nonempty[1][1]
    cover_en_title = nonempty[2][1]
    cover_author = nonempty[3][1]
    cover_school = nonempty[4][1]
    cover_date = nonempty[5][1]
    secrecy = nonempty[6][1]
    info_type = nonempty[7][1]
    info_title = nonempty[8][1]

    for para in doc.paragraphs[:abstract_idx]:
        remove_keep_next(para)
        para.paragraph_format.page_break_before = False

    cover_specs = [
        (cover_type, "黑体", 24, True, 0, 0, 1.0),
        (cover_cn_title, "黑体", 22, True, 12, 0, 1.1),
        (cover_en_title, "Times New Roman", 16, True, 6, 0, 1.1),
        (cover_author, "宋体", 14, False, 28, 0, 1.0),
        (cover_school, "楷体", 18, True, 0, 0, 1.0),
        (cover_date, "宋体", 18, True, 0, 0, 1.0),
        (info_type, "黑体", 18, True, 0, 0, 1.0),
        (info_title, "黑体", 22, True, 0, 0, 1.0),
    ]

    for para, font, size, bold, before, after, line in cover_specs:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        clear_paragraph_indent(para)
        para.paragraph_format.space_before = Pt(before)
        para.paragraph_format.space_after = Pt(after)
        set_line_spacing(para, line)
        format_para_runs(para, east_asia=font, latin="Times New Roman", size=size, bold=bold)

    secrecy.paragraph_format.page_break_before = True
    secrecy.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    clear_paragraph_indent(secrecy)
    secrecy.paragraph_format.space_before = Pt(0)
    secrecy.paragraph_format.space_after = Pt(42)
    set_line_spacing(secrecy, 1.0)
    format_para_runs(secrecy, east_asia="宋体", latin="Times New Roman", size=12, bold=False)

    if doc.tables:
        info_table = doc.tables[0]
        set_no_table_borders(info_table)
        info_table.alignment = 1
        for row in info_table.rows:
            for col_idx, cell in enumerate(row.cells):
                for para in cell.paragraphs:
                    clear_paragraph_indent(para)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    set_line_spacing(para, 1.35)
                    para.alignment = (
                        WD_ALIGN_PARAGRAPH.RIGHT
                        if col_idx == 0
                        else WD_ALIGN_PARAGRAPH.CENTER
                        if col_idx == 1
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=12, bold=False)


def apply_citations(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text or "[" in text and re.search(r"\[\d", text):
            continue
        for marker, cite in CITATION_RULES:
            if text.startswith(marker):
                para.text = citation_insert(text, cite)
                break


def replace_references(doc: Document) -> None:
    ref_anchor = None
    end_anchor = None
    for para in doc.paragraphs:
        text = para.text.strip()
        if text == "参考文献":
            ref_anchor = para
            continue
        if ref_anchor is not None and (text == "攻读学士学位期间取得创新性成果" or is_main_chapter(text)):
            end_anchor = para
            break
    if ref_anchor is None or end_anchor is None:
        raise RuntimeError("Could not locate reference section boundaries.")

    current = ref_anchor._element.getnext()
    while current is not None and current is not end_anchor._element:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt

    current_el = ref_anchor._element
    for item in RECENT_REFERENCES:
        para = insert_paragraph_after(doc, current_el, item, style="Plain Text")
        current_el = para._element


def format_document(doc: Document) -> None:
    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name if para.style else ""
        pf = para.paragraph_format

        is_toc_style = style.startswith("目录") or style.lower().startswith("toc")
        is_real_section_start = style == "Heading 1" and (text in PAGE_START_TITLES or is_main_chapter(text))
        if is_real_section_start:
            ensure_page_break_before(para)
        elif text and not is_toc_style:
            pf.page_break_before = False

        if style == "Heading 1":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(para, 1.2)
            pf.space_before = Pt(0)
            pf.space_after = Pt(12)
            pf.first_line_indent = None
            format_para_runs(para, east_asia="黑体", latin="Times New Roman", size=18, bold=True)
            add_keep_next(para)
        elif style == "Heading 2":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(para, 1.25)
            pf.space_before = Pt(6)
            pf.space_after = Pt(6)
            pf.first_line_indent = None
            format_para_runs(para, east_asia="黑体", latin="Times New Roman", size=15, bold=True)
            add_keep_next(para)
        elif style == "Heading 3":
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(para, 1.2)
            pf.space_before = Pt(6)
            pf.space_after = Pt(3)
            pf.first_line_indent = None
            format_para_runs(para, east_asia="黑体", latin="Times New Roman", size=14, bold=True)
            add_keep_next(para)
        elif is_figure_caption(text) or is_table_caption(text):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_line_spacing(para, 1.0)
            pf.space_before = Pt(3)
            pf.space_after = Pt(6)
            pf.first_line_indent = None
            format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
        elif text.startswith("[") and re.match(r"^\[\d+\]", text):
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            set_line_spacing(para, 1.25)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            pf.left_indent = Cm(0.74)
            pf.first_line_indent = Cm(-0.74)
            format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)
        elif is_toc_style:
            set_line_spacing(para, 1.25)
            pf.space_before = Pt(0)
            pf.space_after = Pt(0)
            format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=12, bold=False)
        elif text and style not in {"Plain Text"}:
            if text not in {"本科毕业论文（设计）", "云平台故障注入工具", "FAULT INJECTION TOOL FOR CLOUD PLATFORMS", "哈尔滨工业大学", "2026年4月", "密级：公开"}:
                pf.line_spacing = 1.5
                set_line_spacing(para, 1.5)
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                if not text.startswith("关键词") and not text.startswith("Keywords"):
                    pf.first_line_indent = Pt(24)
                format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=12, bold=False)

    for table in doc.tables:
        table.alignment = 1
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    set_line_spacing(para, 1.0)
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    format_para_runs(para, east_asia="宋体", latin="Times New Roman", size=10.5, bold=False)

    format_front_matter(doc)


def repair_back_matter(doc: Document) -> None:
    """Restore template-like separation after references."""
    paragraphs = doc.paragraphs
    ref_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "参考文献"), None)
    originality_idx = next((i for i, p in enumerate(paragraphs) if p.text.strip() == "原创性声明和使用权限"), None)
    if ref_idx is None or originality_idx is None:
        return

    # The template has a standalone innovations page between references and the originality statement.
    between = paragraphs[ref_idx + 1 : originality_idx]
    innovation = next((p for p in between if p.text.strip() == "攻读学士学位期间取得创新性成果"), None)
    if innovation is None:
        blank_heading = next((p for p in between if not p.text.strip() and p.style.name == "Heading 1"), None)
        target = blank_heading if blank_heading is not None else paragraphs[originality_idx - 1]
        target.text = "攻读学士学位期间取得创新性成果"
        target.style = "Heading 1"
        innovation = target

    originality_heading_prefix = None
    for para in between:
        text = para.text.strip()
        if not text:
            para.style = doc.styles["Normal"]
            para.paragraph_format.page_break_before = False
        elif text == "哈尔滨工业大学本科毕业论文（设计）":
            originality_heading_prefix = para
            para.style = doc.styles["Normal"]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.page_break_before = True
            para.paragraph_format.first_line_indent = None
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)
            set_line_spacing(para, 1.2)
            format_para_runs(para, east_asia="黑体", latin="Times New Roman", size=18, bold=True)

    for title in POST_REFERENCE_TITLES:
        for para in paragraphs:
            if para.text.strip() == title:
                para.style = "Heading 1"
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.paragraph_format.page_break_before = not (
                    title == "原创性声明和使用权限" and originality_heading_prefix is not None
                )
                para.paragraph_format.first_line_indent = None
                para.paragraph_format.space_before = Pt(0)
                para.paragraph_format.space_after = Pt(12)
                set_line_spacing(para, 1.2)
                format_para_runs(para, east_asia="黑体", latin="Times New Roman", size=18, bold=True)
                break


def set_update_fields_on_open(doc: Document) -> None:
    settings = doc.settings._element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(INPUT, BACKUP)
    shutil.copy2(INPUT, OUTPUT)

    doc = Document(OUTPUT)
    apply_citations(doc)
    replace_references(doc)
    format_document(doc)
    repair_back_matter(doc)
    set_update_fields_on_open(doc)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
