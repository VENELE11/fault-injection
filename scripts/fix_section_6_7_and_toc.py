from __future__ import annotations

import shutil
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "output/doc/云平台故障注入工具-终稿格式精修版.docx"
BACKUP_PATH = ROOT / "output/doc/云平台故障注入工具-终稿格式精修版-修改前备份.docx"


SECTION_6_7_PARAGRAPHS = [
    (
        "本章围绕云平台故障注入工具的实验设计、测试过程与结果分析展开，首先明确了实验环境、"
        "测试对象和评价指标，随后从单层故障注入和跨层级联故障注入两个角度验证了平台能力。"
        "实验结果表明，本文平台能够覆盖宿主机与 KVM 层、Guest OS 层以及 Hadoop、CloudStack "
        "等典型业务层的多类故障场景，并能够通过统一控制面完成参数配置、任务触发、日志回收和"
        "状态对比，为云平台健壮性测试提供较完整的实验支撑。"
    ),
    (
        "从单层实验结果看，平台能够较稳定地触发 CPU 资源争抢、内存压力、网络异常、进程干预、"
        "KVM 性能扰动以及 CloudStack 管理链路异常等故障，并通过前后状态对比呈现被测系统在资源"
        "隔离、服务保持和恢复处理方面的差异。从跨层级联实验结果看，底层虚拟化扰动、Guest OS "
        "资源异常和上层管理服务异常之间存在一定传播关系，单一层次测试难以完整暴露这类问题，"
        "因此多层次联合注入对于发现云平台潜在健壮性风险具有必要性。"
    ),
    (
        "同时，本章实验也说明本文平台仍存在进一步完善空间，例如结果分析仍以日志对比和现象"
        "归纳为主，自动化指标统计、可视化趋势分析和更大规模场景覆盖还有待增强。总体而言，"
        "本章验证了本文故障注入工具在功能覆盖、执行闭环和实验可复现性方面的有效性，也为后续"
        "结论部分总结全文工作、分析不足并提出改进方向奠定了依据。"
    ),
]


def set_chinese_font(run, size: float = 12.0) -> None:
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)


def set_body_format(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.first_line_indent = Cm(0.74)
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(0)
    for run in paragraph.runs:
        set_chinese_font(run, 12.0)


def insert_after(paragraph, text: str, style_name: str):
    new_p = paragraph._p.__class__()
    paragraph._p.addnext(new_p)
    new_para = paragraph._parent.paragraphs[-1]
    # python-docx does not return the inserted wrapper from addnext; locate it by XML identity.
    for candidate in paragraph._parent.paragraphs:
        if candidate._p is new_p:
            new_para = candidate
            break
    new_para.style = style_name
    new_para.add_run(text)
    set_body_format(new_para)
    return new_para


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)
    paragraph._p = paragraph._element = None


def main() -> None:
    if not BACKUP_PATH.exists():
        shutil.copy2(DOC_PATH, BACKUP_PATH)

    doc = Document(DOC_PATH)

    # Remove stale TOC item only. The originality/use-permission text remains unchanged.
    for paragraph in list(doc.paragraphs):
        if paragraph.text.strip().startswith("攻读学士学位期间取得创新性成果"):
            remove_paragraph(paragraph)
            break

    heading = None
    for paragraph in doc.paragraphs:
        if paragraph.text.strip() == "6.7 本章小结":
            heading = paragraph
            break
    if heading is None:
        raise RuntimeError("Could not find section heading: 6.7 本章小结")

    # Avoid duplicate insertion if the script is re-run.
    next_text = ""
    paragraphs = doc.paragraphs
    for idx, paragraph in enumerate(paragraphs):
        if paragraph is heading and idx + 1 < len(paragraphs):
            next_text = paragraphs[idx + 1].text.strip()
            break
    if next_text == "结  论":
        anchor = heading
        style_name = "正文首行缩进"
        for text in SECTION_6_7_PARAGRAPHS:
            anchor = insert_after(anchor, text, style_name)

    doc.save(DOC_PATH)
    print(DOC_PATH)
    print(BACKUP_PATH)


if __name__ == "__main__":
    main()
