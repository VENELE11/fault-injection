from __future__ import annotations

import re
import shutil
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from docx.text.paragraph import Paragraph


ROOT = Path("/Users/venele/Downloads/fault-injection")
TEMPLATE = ROOT / "论文模板-工作版.docx"
SOURCE = ROOT / "output/doc/云平台故障注入工具-论文统一版.docx"
OUTPUT = ROOT / "output/doc/云平台故障注入工具-模板填充版.docx"
ASSET_DIR = ROOT / "output/doc/assets/unified"


TITLE_CN = "云平台故障注入工具"
TITLE_EN = "FAULT INJECTION TOOL FOR CLOUD PLATFORMS"
DATE_TEXT = "2026年4月"

FIGURE_PATHS = {
    "图2-1": ASSET_DIR / "fig2_1.png",
    "图2-2": ASSET_DIR / "fig2_2.png",
    "图2-3": ASSET_DIR / "fig2_3.png",
    "图2-4": ASSET_DIR / "fig2_4.png",
    "图3-1": ASSET_DIR / "fig3_1.png",
    "图3-2": ASSET_DIR / "fig3_2.png",
    "图4-1": ASSET_DIR / "fig4_1.png",
    "图4-2": ASSET_DIR / "fig4_2.png",
    "图4-3": ASSET_DIR / "fig4_3.png",
    "图4-4": ASSET_DIR / "fig4_4.png",
    "图5-1": ASSET_DIR / "fig5_1.png",
    "图5-2": ASSET_DIR / "fig5_2.png",
}

SOURCE_FIGURE_INSERTS = {
    51: ("图2-1", "图2-1 云平台分层架构示意图"),
    73: ("图2-2", "图2-2 故障注入实验环境的物理与逻辑结构"),
    78: ("图2-3", "图2-3 轻量级虚拟机集群网络拓扑"),
    82: ("图2-4", "图2-4 本课题多层次故障注入总体结构"),
    101: ("图3-1", "图3-1 云平台故障注入工具总体架构图"),
    140: ("图3-2", "图3-2 故障注入任务执行流程图"),
    166: ("图4-1", "图4-1 缺页异常拦截与异常传递过程示意"),
    191: ("图4-2", "图4-2 基于 POSIX 信号的进程状态控制实现示意"),
    204: ("图4-3", "图4-3 基于 virsh 的 CloudStack SystemVM 干预实现示意"),
    210: ("图4-4", "图4-4 故障注入任务执行生命周期示意图"),
    224: ("图5-1", "图5-1 平台控制面内部结构图"),
    255: ("图5-2", "图5-2 控制面与多节点执行环境集成示意图"),
}

FIGURE_CAPTION_PARAS = {
    51,
    73,
    82,
    101,
    140,
    166,
    191,
    204,
    210,
    224,
    255,
    283,
    288,
    292,
    298,
}

TABLE_CAPTION_PARAS = {
    268: 0,
    276: 1,
}

CH_ABS_EN = [
    "With the rapid expansion of cloud computing platforms, the coupling among the virtualization layer, the operating system runtime, and upper-layer distributed applications has become increasingly complex. Faults such as single-point failures, resource exhaustion, network anomalies, and control-path disruption can easily propagate across layers and be amplified, eventually affecting platform availability and robustness. To expose such risks before deployment, it is necessary to build a fault injection tool for cloud platforms and carry out systematic robustness testing under representative fault scenarios.",
    "This thesis designs and implements a multi-layer fault injection tool for cloud platforms. First, an experimental environment based on ARM64 architecture, KVM virtualization, and a lightweight virtual machine cluster is constructed to support low-level fault injection and cross-node experiments. Second, for the IaaS layer, a kernel-level fault injection mechanism is implemented by using dynamic probing techniques such as Kprobes to interfere with key paths including virtual CPU state, memory mapping, TLB refresh, virtual file system I/O, and KVM state handling. Third, for the Guest OS layer, application-level injectors are developed for resource exhaustion, network anomalies, process suspension, and memory tampering. Scenario-oriented injection methods are further designed for Hadoop and CloudStack to support comprehensive fault-tolerance testing of typical cloud applications. Finally, experiments under single-layer and cross-layer cascading fault conditions are conducted to verify the effectiveness of the platform.",
    "The results show that the proposed tool can simulate common computing, memory, storage, network, and management faults in cloud platforms with good controllability. It can effectively trigger abnormal behaviors and expose weaknesses in resource isolation, fault detection, recovery mechanisms, and service continuity. This work provides practical tool support and an experimental basis for robustness testing and fault-tolerance evaluation of cloud platforms.",
]
ABS_EN_KEYWORDS = "Keywords: cloud platform; fault injection; KVM; cloud computing; robustness testing; fault-tolerance evaluation"


def remove_paragraph(para: Paragraph) -> None:
    el = para._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def remove_between(start_para: Paragraph, end_para: Paragraph) -> None:
    current = start_para._element.getnext()
    while current is not None and current is not end_para._element:
        nxt = current.getnext()
        current.getparent().remove(current)
        current = nxt


def remove_before(first_keep: Paragraph) -> None:
    current = first_keep._element.getprevious()
    while current is not None:
        prev = current.getprevious()
        current.getparent().remove(current)
        current = prev


def remove_after(last_keep: Paragraph) -> None:
    current = last_keep._element.getnext()
    while current is not None:
        nxt = current.getnext()
        if current.tag.endswith("sectPr"):
            break
        current.getparent().remove(current)
        current = nxt


def insert_paragraph_after(
    doc: Document,
    current_el,
    text: str = "",
    style: str | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
    page_break_before: bool = False,
) -> Paragraph:
    new_p = deepcopy(doc.paragraphs[0]._element)
    for child in list(new_p):
        new_p.remove(child)
    current_el.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    para.text = text
    if style:
        para.style = style
    if align is not None:
        para.alignment = align
    if page_break_before:
        para.paragraph_format.page_break_before = True
    return para


def insert_picture_after(doc: Document, current_el, image_path: Path, width_cm: float = 14.2):
    para = insert_paragraph_after(doc, current_el, style="Normal", align=WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run()
    run.add_picture(str(image_path), width=Cm(width_cm))
    return para._element


def insert_table_after(doc: Document, current_el, src_table) -> object:
    new_tbl = deepcopy(src_table._tbl)
    current_el.addnext(new_tbl)
    return new_tbl


def format_main_heading(text: str) -> str:
    if text == "第1章 绪论":
        return "第1章  绪  论"
    return text.replace("第", "第", 1).replace("章 ", "章  ", 1)


def classify_style(text: str) -> str:
    if re.match(r"^第[1-9]\d*章(\s|　)", text):
        return "Heading 1"
    if re.match(r"^\d+\.\d+\.\d+", text):
        return "Heading 3"
    if re.match(r"^\d+\.\d+", text):
        return "Heading 2"
    return "正文首行缩进"


def replace_cover(template_doc: Document) -> None:
    for old, new in [
        ("本科毕业论文（设计）书写范例", "本科毕业论文（设计）"),
        ("（理工类）", "（理工类）"),
        ("局部多孔质气体静压轴承关键技术的研究", TITLE_CN),
        ("RESEARCH ON KEY TECHNOLOGIES OF PARTIAL POROUS EXTERNALLY PRESSURIZED GAS BEARING", TITLE_EN),
    ]:
        for para in template_doc.paragraphs:
            if para.text.strip() == old:
                para.text = new

    template_doc.paragraphs[42].text = "哈尔滨工业大学"
    template_doc.paragraphs[43].text = DATE_TEXT
    template_doc.paragraphs[45].text = "密级：公开"
    template_doc.paragraphs[49].paragraph_format.page_break_before = True

    cover_table = template_doc.tables[0]
    cover_table.cell(0, 2).text = "□□□"
    cover_table.cell(1, 2).text = "□□□□□□□□□□"
    cover_table.cell(2, 2).text = "□□□教授"
    cover_table.cell(3, 2).text = "□□□□□□□□"
    cover_table.cell(4, 2).text = "□□□□□□"
    cover_table.cell(5, 2).text = "20□□年□月"
    cover_table.cell(6, 2).text = "哈尔滨工业大学"

    statement = template_doc.paragraphs[277].text
    template_doc.paragraphs[277].text = re.sub(r"《\s*》", f"《{TITLE_CN}》", statement)

    format_cover(template_doc)


def set_run_font(run, font_name: str | None = None, size_pt: float | None = None, bold: bool | None = None) -> None:
    if font_name:
        run.font.name = font_name
        run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    if size_pt is not None:
        run.font.size = Pt(size_pt)
    if bold is not None:
        run.font.bold = bold


def replace_para_with_single_run(
    para: Paragraph,
    text: str,
    *,
    font_name: str | None = None,
    size_pt: float | None = None,
    bold: bool | None = None,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    p = para._element
    for child in list(p):
        if child.tag != qn("w:pPr"):
            p.remove(child)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run, font_name=font_name, size_pt=size_pt, bold=bold)


def format_cover(doc: Document) -> None:
    replace_para_with_single_run(doc.paragraphs[25], TITLE_CN, font_name="黑体", size_pt=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_para_with_single_run(doc.paragraphs[28], TITLE_EN, font_name="Times New Roman", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_para_with_single_run(doc.paragraphs[42], "哈尔滨工业大学", font_name="楷体", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_para_with_single_run(doc.paragraphs[43], DATE_TEXT, font_name="宋体", size_pt=18, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_para_with_single_run(doc.paragraphs[45], "密级：公开", font_name="宋体", size_pt=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    replace_para_with_single_run(doc.paragraphs[52], TITLE_CN, font_name="黑体", size_pt=22, align=WD_ALIGN_PARAGRAPH.CENTER)

    cover_table = doc.tables[0]
    for row_idx in range(7):
        cell_para = cover_table.cell(row_idx, 2).paragraphs[0]
        text = cell_para.text
        replace_para_with_single_run(cell_para, text, font_name="宋体", size_pt=14)


def remove_instruction_textboxes(doc: Document) -> None:
    to_remove = []
    for txbx in doc._element.xpath('.//*[local-name()="txbx"]'):
        txt = "".join(t.text or "" for t in txbx.xpath('.//*[local-name()="t"]'))
        if "冒号左侧用黑体" not in txt:
            continue
        node = txbx
        for _ in range(8):
            node = node.getparent()
        to_remove.append(node)

    for node in to_remove:
        parent = node.getparent()
        if parent is not None:
            parent.remove(node)


def clean_template(template_doc: Document) -> dict[str, Paragraph]:
    removable_notes = [template_doc.paragraphs[idx] for idx in [23, 24, 26, 27, 29, 30, 33, 34, 44, 50, 51, 53, 54]]
    anchors = {
        "cover_start": template_doc.paragraphs[22],
        "abstract": template_doc.paragraphs[62],
        "abstract_en": template_doc.paragraphs[75],
        "toc": template_doc.paragraphs[87],
        "chapter1": template_doc.paragraphs[117],
        "conclusion": template_doc.paragraphs[227],
        "references": template_doc.paragraphs[239],
        "innovation": template_doc.paragraphs[258],
        "statement": template_doc.paragraphs[274],
        "ack": template_doc.paragraphs[292],
    }

    remove_before(anchors["cover_start"])

    for para in removable_notes:
        remove_paragraph(para)

    remove_between(anchors["abstract"], anchors["abstract_en"])
    remove_between(anchors["abstract_en"], anchors["toc"])
    remove_between(anchors["chapter1"], anchors["conclusion"])
    remove_between(anchors["conclusion"], anchors["references"])
    remove_between(anchors["references"], anchors["innovation"])
    remove_between(anchors["innovation"], anchors["statement"])
    remove_after(anchors["ack"])

    return anchors


def fill_abstract(template_doc: Document, anchors: dict[str, Paragraph], source_doc: Document) -> None:
    abs_anchor = anchors["abstract"]
    abs_en_anchor = anchors["abstract_en"]

    abs_anchor.text = "摘  要"
    current = abs_anchor._element
    for idx in [3, 4, 5]:
        para = insert_paragraph_after(template_doc, current, source_doc.paragraphs[idx].text.strip(), style="正文首行缩进")
        current = para._element
    para = insert_paragraph_after(
        template_doc,
        current,
        source_doc.paragraphs[6].text.strip().replace("关键词：", "关键词："),
        style="Normal",
    )
    current = para._element

    abs_en_anchor.text = "Abstract"
    current = abs_en_anchor._element
    for text in CH_ABS_EN:
        para = insert_paragraph_after(template_doc, current, text, style="Normal")
        current = para._element
    insert_paragraph_after(template_doc, current, ABS_EN_KEYWORDS, style="Normal")


def fill_main_body(template_doc: Document, anchors: dict[str, Paragraph], source_doc: Document) -> None:
    anchors["chapter1"].text = format_main_heading(source_doc.paragraphs[11].text.strip())
    current = anchors["chapter1"]._element

    for idx in range(12, 309):
        text = source_doc.paragraphs[idx].text.strip()
        if not text:
            continue

        if idx in SOURCE_FIGURE_INSERTS:
            fig_no, caption = SOURCE_FIGURE_INSERTS[idx]
            current = insert_picture_after(template_doc, current, FIGURE_PATHS[fig_no])
            caption_para = insert_paragraph_after(
                template_doc,
                current,
                caption,
                style="Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            current = caption_para._element
            if idx != 78:
                continue

        if idx in TABLE_CAPTION_PARAS:
            caption_para = insert_paragraph_after(
                template_doc,
                current,
                text,
                style="Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            current = caption_para._element
            src_table = source_doc.tables[TABLE_CAPTION_PARAS[idx]]
            current = insert_table_after(template_doc, current, src_table)
            continue

        if idx in FIGURE_CAPTION_PARAS:
            para = insert_paragraph_after(
                template_doc,
                current,
                text,
                style="Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER,
            )
            current = para._element
            continue

        style = classify_style(text)
        para = insert_paragraph_after(
            template_doc,
            current,
            format_main_heading(text) if style == "Heading 1" else text,
            style=style,
            page_break_before=(style == "Heading 1"),
        )
        current = para._element


def fill_conclusion(template_doc: Document, anchors: dict[str, Paragraph], source_doc: Document) -> None:
    anchors["conclusion"].text = "结  论"
    current = anchors["conclusion"]._element

    for idx in [313, 314, 315, 316, 317]:
        para = insert_paragraph_after(template_doc, current, source_doc.paragraphs[idx].text.strip(), style="正文首行缩进")
        current = para._element

    for idx in [319, 320, 321, 322, 324, 325, 326, 327, 328]:
        para = insert_paragraph_after(template_doc, current, source_doc.paragraphs[idx].text.strip(), style="正文首行缩进")
        current = para._element


def fill_references(template_doc: Document, anchors: dict[str, Paragraph], source_doc: Document) -> None:
    anchors["references"].text = "参考文献"
    current = anchors["references"]._element
    for idx in range(331, 352):
        para = insert_paragraph_after(template_doc, current, source_doc.paragraphs[idx].text.strip(), style="Plain Text")
        current = para._element


def fill_innovation_and_ack(template_doc: Document, anchors: dict[str, Paragraph]) -> None:
    current = anchors["innovation"]._element
    para = insert_paragraph_after(template_doc, current, "无。", style="正文首行缩进")
    current = para._element

    ack_current = anchors["ack"]._element
    ack_paras = [
        "在本课题研究与论文撰写过程中，得到了指导教师、实验室老师和同学们的帮助与支持，在此谨致以诚挚的感谢。",
        "感谢学院和学校提供的实验条件与学习环境，感谢相关开源社区、技术资料和参考文献为本课题提供的研究基础。",
        "同时感谢家人和朋友在毕业设计期间给予的理解、鼓励与陪伴。",
    ]
    for i, text in enumerate(ack_paras):
        style = "Body Text Indent" if i == 0 else "Normal Indent"
        para = insert_paragraph_after(template_doc, ack_current, text, style=style)
        ack_current = para._element


def adjust_tables(doc: Document) -> None:
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER


def enforce_page_starts(doc: Document) -> None:
    fixed_titles = {
        "摘  要",
        "Abstract",
        "目  录",
        "结  论",
        "参考文献",
        "攻读学士学位期间取得创新性成果",
        "原创性声明和使用权限",
        "致  谢",
    }
    chapter_re = re.compile(r"^第[1-9]\d*章")
    for para in doc.paragraphs:
        text = para.text.strip()
        if text in fixed_titles or chapter_re.match(text):
            para.paragraph_format.page_break_before = True


def main() -> None:
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(TEMPLATE, OUTPUT)

    template_doc = Document(OUTPUT)
    source_doc = Document(SOURCE)

    replace_cover(template_doc)
    anchors = clean_template(template_doc)
    fill_abstract(template_doc, anchors, source_doc)
    fill_main_body(template_doc, anchors, source_doc)
    fill_conclusion(template_doc, anchors, source_doc)
    fill_references(template_doc, anchors, source_doc)
    fill_innovation_and_ack(template_doc, anchors)
    remove_instruction_textboxes(template_doc)
    enforce_page_starts(template_doc)
    adjust_tables(template_doc)

    template_doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
