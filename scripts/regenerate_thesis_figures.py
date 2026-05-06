from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
import math
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/venele/Downloads/fault-injection")
OUT_DIR = ROOT / "output/doc/assets/unified"
PREVIEW_PATH = OUT_DIR / "_preview_grid.png"

DOCS = [
    ROOT / "output/doc/云平台故障注入工具-论文统一版.docx",
    ROOT / "output/doc/云平台故障注入工具-论文初稿.docx",
]

FONT_MEDIUM = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_LIGHT = "/System/Library/Fonts/STHeiti Light.ttc"

COLORS = {
    "ink": "#4B5563",
    "text": "#2E3440",
    "subtext": "#6B7280",
    "blue": "#DDEAF6",
    "green": "#DDEBDC",
    "purple": "#E8DFF3",
    "sand": "#F6EAD2",
    "mint": "#E1EEE4",
    "paper": "#FFFFFF",
}


@dataclass
class Card:
    title: str
    body: str = ""
    fill: str = COLORS["paper"]
    outline: str = COLORS["ink"]
    title_size: int = 34
    body_size: int = 22
    pad_x: int = 30
    pad_y: int = 22
    gap: int = 10
    radius: int = 24
    min_w: int = 260
    max_w: int = 520
    align: str = "center"
    title_color: str = COLORS["text"]
    body_color: str = COLORS["subtext"]
    x: int = 0
    y: int = 0
    w: int = 0
    h: int = 0

    def layout(self, measure: "Measure", content_w: int | None = None) -> "Card":
        title_font = measure.font(FONT_MEDIUM, self.title_size)
        body_font = measure.font(FONT_LIGHT, self.body_size)

        if content_w is None:
            content_w = self.max_w - 2 * self.pad_x
        title_lines = wrap_text(measure, self.title, title_font, content_w)
        body_lines = wrap_text(measure, self.body, body_font, content_w) if self.body else []

        title_w, title_h = text_block_size(measure, title_lines, title_font, 8)
        body_w, body_h = text_block_size(measure, body_lines, body_font, 6)
        inner_w = max(title_w, body_w)
        self.w = max(self.min_w, min(self.max_w, inner_w + self.pad_x * 2))

        final_content_w = self.w - self.pad_x * 2
        title_lines = wrap_text(measure, self.title, title_font, final_content_w)
        body_lines = wrap_text(measure, self.body, body_font, final_content_w) if self.body else []
        title_w, title_h = text_block_size(measure, title_lines, title_font, 8)
        body_w, body_h = text_block_size(measure, body_lines, body_font, 6)

        used_gap = self.gap if body_lines else 0
        self.h = self.pad_y * 2 + title_h + body_h + used_gap
        self._title_lines = title_lines
        self._body_lines = body_lines
        self._title_font = title_font
        self._body_font = body_font
        self._title_h = title_h
        self._used_gap = used_gap
        return self

    @property
    def bounds(self) -> tuple[int, int, int, int]:
        return (self.x, self.y, self.x + self.w, self.y + self.h)


@dataclass
class GroupBox:
    title: str
    fill: str
    items: list[Card]
    pad_l: int = 40
    pad_r: int = 40
    pad_t: int = 100
    pad_b: int = 40
    title_size: int = 34
    x1: int = 0
    y1: int = 0
    x2: int = 0
    y2: int = 0

    def layout(self, measure: "Measure") -> "GroupBox":
        bounds = union_bounds([item.bounds for item in self.items])
        title_font = measure.font(FONT_MEDIUM, self.title_size)
        title_bbox = measure.text_bbox((0, 0), self.title, title_font)
        title_w = title_bbox[2] - title_bbox[0]
        title_h = title_bbox[3] - title_bbox[1]

        self.x1 = bounds[0] - self.pad_l
        self.y1 = bounds[1] - self.pad_t
        self.x2 = max(bounds[2] + self.pad_r, bounds[0] + title_w + self.pad_l * 2)
        self.y2 = bounds[3] + self.pad_b
        self._title_font = title_font
        self._title_h = title_h
        return self

    @property
    def bounds(self) -> tuple[int, int, int, int]:
        return (self.x1, self.y1, self.x2, self.y2)


class Measure:
    def __init__(self) -> None:
        self.image = Image.new("RGB", (32, 32), "white")
        self.draw = ImageDraw.Draw(self.image)
        self.font_cache: dict[tuple[str, int], ImageFont.FreeTypeFont] = {}

    def font(self, path: str, size: int) -> ImageFont.FreeTypeFont:
        key = (path, size)
        if key not in self.font_cache:
            self.font_cache[key] = ImageFont.truetype(path, size)
        return self.font_cache[key]

    def text_bbox(self, pos: tuple[int, int], text: str, font: ImageFont.FreeTypeFont) -> tuple[int, int, int, int]:
        return self.draw.textbbox(pos, text, font=font)


def wrap_text(measure: Measure, text: str, font: ImageFont.FreeTypeFont, max_width: int) -> list[str]:
    if not text:
        return []

    lines: list[str] = []
    for paragraph in text.split("\n"):
        if not paragraph:
            lines.append("")
            continue
        current = ""
        for ch in paragraph:
            test = current + ch
            bbox = measure.text_bbox((0, 0), test, font)
            if bbox[2] - bbox[0] <= max_width or not current:
                current = test
            else:
                lines.append(current)
                current = ch
        if current:
            lines.append(current)
    return lines


def text_block_size(
    measure: Measure,
    lines: list[str],
    font: ImageFont.FreeTypeFont,
    spacing: int,
) -> tuple[int, int]:
    if not lines:
        return 0, 0
    widths = []
    heights = []
    for line in lines:
        bbox = measure.text_bbox((0, 0), line or " ", font)
        widths.append(bbox[2] - bbox[0])
        heights.append(bbox[3] - bbox[1])
    total_h = sum(heights) + spacing * (len(lines) - 1)
    return max(widths), total_h


def rounded(draw: ImageDraw.ImageDraw, bounds: tuple[int, int, int, int], fill: str, outline: str, radius: int = 24, width: int = 2) -> None:
    draw.rounded_rectangle(bounds, radius=radius, fill=fill, outline=outline, width=width)


def draw_card(draw: ImageDraw.ImageDraw, card: Card) -> None:
    rounded(draw, card.bounds, card.fill, card.outline, radius=card.radius)
    cx = card.x + card.w // 2
    y = card.y + card.pad_y

    title_font = card._title_font
    body_font = card._body_font

    for line in card._title_lines:
        bbox = draw.textbbox((0, 0), line or " ", font=title_font)
        lw = bbox[2] - bbox[0]
        lh = bbox[3] - bbox[1]
        tx = card.x + card.pad_x if card.align == "left" else cx - lw // 2
        draw.text((tx, y), line, font=title_font, fill=card.title_color)
        y += lh + 8

    if card._body_lines:
        y += card.gap - 8
        for line in card._body_lines:
            bbox = draw.textbbox((0, 0), line or " ", font=body_font)
            lw = bbox[2] - bbox[0]
            lh = bbox[3] - bbox[1]
            tx = card.x + card.pad_x if card.align == "left" else cx - lw // 2
            draw.text((tx, y), line, font=body_font, fill=card.body_color)
            y += lh + 6


def draw_group(draw: ImageDraw.ImageDraw, group: GroupBox) -> None:
    rounded(draw, group.bounds, group.fill, COLORS["ink"], radius=30)
    draw.text((group.x1 + group.pad_l, group.y1 + 34), group.title, font=group._title_font, fill=COLORS["text"])


def union_bounds(bounds: list[tuple[int, int, int, int]]) -> tuple[int, int, int, int]:
    x1 = min(b[0] for b in bounds)
    y1 = min(b[1] for b in bounds)
    x2 = max(b[2] for b in bounds)
    y2 = max(b[3] for b in bounds)
    return (x1, y1, x2, y2)


def point_from(card: Card | GroupBox, anchor: str) -> tuple[int, int]:
    if isinstance(card, GroupBox):
        x1, y1, x2, y2 = card.bounds
    else:
        x1, y1, x2, y2 = card.bounds
    mapping = {
        "top": ((x1 + x2) // 2, y1),
        "bottom": ((x1 + x2) // 2, y2),
        "left": (x1, (y1 + y2) // 2),
        "right": (x2, (y1 + y2) // 2),
        "top_left": (x1, y1),
        "top_right": (x2, y1),
        "bottom_left": (x1, y2),
        "bottom_right": (x2, y2),
    }
    return mapping[anchor]


def draw_arrow(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 4) -> None:
    draw.line([start, end], fill=COLORS["ink"], width=width)
    add_arrow_head(draw, start, end, width)


def draw_poly_arrow(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], width: int = 4) -> None:
    draw.line(points, fill=COLORS["ink"], width=width)
    add_arrow_head(draw, points[-2], points[-1], width)


def add_arrow_head(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], width: int = 4) -> None:
    dx = end[0] - start[0]
    dy = end[1] - start[1]
    length = math.hypot(dx, dy) or 1
    ux = dx / length
    uy = dy / length
    size = 16 + width
    left = (
        end[0] - ux * size - uy * size * 0.5,
        end[1] - uy * size + ux * size * 0.5,
    )
    right = (
        end[0] - ux * size + uy * size * 0.5,
        end[1] - uy * size - ux * size * 0.5,
    )
    draw.polygon([end, left, right], fill=COLORS["ink"])


def crop_and_save(image: Image.Image, bounds: list[tuple[int, int, int, int]], target: Path, margin: int = 50) -> None:
    bbox = union_bounds(bounds)
    x1 = max(0, bbox[0] - margin)
    y1 = max(0, bbox[1] - margin)
    x2 = min(image.width, bbox[2] + margin)
    y2 = min(image.height, bbox[3] + margin)
    cropped = image.crop((x1, y1, x2, y2))
    cropped.save(target, format="PNG")


def canvas() -> tuple[Image.Image, ImageDraw.ImageDraw, list[tuple[int, int, int, int]]]:
    image = Image.new("RGB", (2200, 1800), "white")
    draw = ImageDraw.Draw(image)
    return image, draw, []


def add_bounds(bounds: list[tuple[int, int, int, int]], *items: Card | GroupBox | tuple[int, int, int, int]) -> None:
    for item in items:
        bounds.append(item if isinstance(item, tuple) else item.bounds)


def h_center(cards: list[Card], y: int, gap: int, start_x: int) -> None:
    x = start_x
    for card in cards:
        card.x = x
        card.y = y
        x += card.w + gap


def fig2_1(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    widths = 1260
    start_x = 420
    y = 120
    layers = [
        Card("用户接入层", "管理员 / 测试人员 / 用户请求", fill=COLORS["blue"], min_w=widths, max_w=widths),
        Card("云原生编排层", "k3s 控制面 / Kubernetes API / 调度与编排", fill=COLORS["sand"], min_w=widths, max_w=widths),
        Card("资源与虚拟化层", "ARM64 宿主机 / KVM / 虚拟网络 / 存储资源", fill=COLORS["green"], min_w=widths, max_w=widths),
        Card("虚拟机与系统层", "Guest OS / 轻量级虚拟机 / 系统服务", fill=COLORS["purple"], min_w=widths, max_w=widths),
        Card("应用与混沌实验层", "目标 Pod / Service / Chaos Mesh 实验对象", fill=COLORS["blue"], min_w=widths, max_w=widths),
    ]
    for layer in layers:
        layer.layout(measure, content_w=widths - 80)
        layer.x = start_x
        layer.y = y
        draw_card(draw, layer)
        add_bounds(bounds, layer)
        y += layer.h + 48

    for first, second in zip(layers, layers[1:]):
        draw_arrow(draw, point_from(first, "bottom"), point_from(second, "top"))

    callout = Card(
        "故障注入观察点",
        "可在资源层、系统层和业务层\n分别实施注入与结果观测",
        fill=COLORS["mint"],
        min_w=360,
        max_w=420,
        title_size=30,
        body_size=20,
    ).layout(measure)
    callout.x = 60
    callout.y = 405
    draw_card(draw, callout)
    add_bounds(bounds, callout)

    src = point_from(callout, "right")
    spine_x = 385
    target_layers = [layers[2], layers[3], layers[4]]
    target_ys = [layer.y + layer.h // 2 for layer in target_layers]
    draw.line([src, (spine_x, src[1])], fill=COLORS["ink"], width=3)
    draw.line([(spine_x, min(src[1], min(target_ys))), (spine_x, max(target_ys))], fill=COLORS["ink"], width=3)
    for layer, target_y in zip(target_layers, target_ys):
        draw_arrow(draw, (spine_x, target_y), (layer.x - 10, target_y), width=3)
    crop_and_save(image, bounds, target)


def fig2_2(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    right_cards = [
        Card("KVM 虚拟化环境", "虚拟 CPU / 内存 / VFS / 网络路径", fill=COLORS["paper"], min_w=420, max_w=460).layout(measure),
        Card("轻量级虚拟机集群", "控制节点 / 工作节点 / Guest OS", fill=COLORS["paper"], min_w=420, max_w=460).layout(measure),
        Card("云原生应用场景", "k3s 工作负载 / Chaos Mesh / 事件状态", fill=COLORS["paper"], min_w=420, max_w=460).layout(measure),
    ]
    for idx, card in enumerate(right_cards):
        card.x = 980
        card.y = 250 + idx * 270

    left_cards = [
        Card("ARM64 宿主机 1", "KVM / Libvirt / 注入模块", fill=COLORS["paper"], min_w=360, max_w=380).layout(measure),
        Card("ARM64 宿主机 2", "备用节点 / 管理节点", fill=COLORS["paper"], min_w=360, max_w=380).layout(measure),
        Card("实验管理网络", "SSH 管理通道 / 日志回收链路", fill=COLORS["paper"], min_w=360, max_w=380).layout(measure),
    ]
    for left, right in zip(left_cards, right_cards):
        left.x = 150
        left.y = right.y

    left_group = GroupBox("物理资源层", COLORS["sand"], left_cards, pad_t=110).layout(measure)
    right_group = GroupBox("逻辑实验层", COLORS["green"], right_cards, pad_t=110).layout(measure)

    for group in (left_group, right_group):
        draw_group(draw, group)
        add_bounds(bounds, group)
    for card in left_cards + right_cards:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(left_cards[0], "right"), point_from(right_cards[0], "left"))
    draw_arrow(draw, point_from(left_cards[1], "right"), point_from(right_cards[1], "left"))
    draw_arrow(draw, point_from(left_cards[2], "right"), point_from(right_cards[2], "left"))
    draw_arrow(draw, point_from(right_cards[0], "bottom"), point_from(right_cards[1], "top"))
    draw_arrow(draw, point_from(right_cards[1], "bottom"), point_from(right_cards[2], "top"))
    crop_and_save(image, bounds, target)


def fig2_3(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    top = Card("控制节点", "Web 控制台 / SSH 指令下发", fill=COLORS["blue"], min_w=500, max_w=560).layout(measure)
    top.x, top.y = 840, 110

    center = Card("宿主机网桥 br0", "管理网络与实验网络汇聚", fill=COLORS["sand"], min_w=540, max_w=600).layout(measure)
    center.x, center.y = 820, 380

    left = Card("管理网络", "节点接入 / 日志回传", fill=COLORS["mint"], min_w=300, max_w=340).layout(measure)
    left.x, left.y = 230, 420

    vms = [
        Card("VM 控制节点", "k3s server / API Server", fill=COLORS["paper"], min_w=330, max_w=360).layout(measure),
        Card("VM 工作节点 1", "k3s agent / 目标 Pod", fill=COLORS["paper"], min_w=330, max_w=360).layout(measure),
        Card("VM 工作节点 2", "k3s agent / 目标 Pod", fill=COLORS["paper"], min_w=330, max_w=360).layout(measure),
    ]
    for idx, card in enumerate(vms):
        card.x = 360 + idx * 430
        card.y = 760

    system_vm = Card("Chaos Mesh 组件", "Controller / Daemon / Dashboard", fill=COLORS["purple"], min_w=420, max_w=460).layout(measure)
    system_vm.x, system_vm.y = 1470, 430

    all_cards = [top, center, left, system_vm] + vms
    for card in all_cards:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(top, "bottom"), point_from(center, "top"))
    draw_arrow(draw, point_from(left, "right"), point_from(center, "left"))
    for vm in vms:
        draw_arrow(draw, point_from(center, "bottom"), point_from(vm, "top"))
    draw_arrow(draw, point_from(center, "right"), point_from(system_vm, "left"))
    crop_and_save(image, bounds, target)


def fig2_4(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    control = Card("统一控制面", "任务编排 / 参数配置 / 结果回收", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure)
    control.x, control.y = 690, 110

    host = Card("IaaS 层注入器", "Kprobes / KVM / VFS / 缺页路径", fill=COLORS["sand"], min_w=360, max_w=400).layout(measure)
    guest = Card("Guest OS 层注入器", "CPU 争抢 / 内存耗尽 / 网络异常 / 进程控制", fill=COLORS["green"], min_w=400, max_w=430).layout(measure)
    app = Card("场景化注入器", "k3s / Chaos Mesh 故障场景", fill=COLORS["purple"], min_w=360, max_w=390).layout(measure)
    host.x, host.y = 200, 430
    guest.x, guest.y = 820, 430
    app.x, app.y = 1430, 430

    observe = Card("监测与结果分析", "状态监控 / 日志采集 / 异常恢复评估", fill=COLORS["mint"], min_w=860, max_w=920).layout(measure)
    observe.x, observe.y = 640, 860

    for card in [control, host, guest, app, observe]:
        draw_card(draw, card)
        add_bounds(bounds, card)

    for target_card in (host, guest, app):
        draw_arrow(draw, point_from(control, "bottom"), point_from(target_card, "top"))
        draw_arrow(draw, point_from(target_card, "bottom"), point_from(observe, "top"))
    crop_and_save(image, bounds, target)


def fig3_1(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    ctrl_cards = [
        Card("任务解析器", "读取故障类型\n目标节点与参数", min_w=320, max_w=340).layout(measure),
        Card("故障模型管理器", "统一描述 CPU / 内存 / 网络\n进程 / 控制面故障", min_w=380, max_w=420).layout(measure),
        Card("调度与分发器", "组织执行顺序\n远程下发任务", min_w=320, max_w=340).layout(measure),
        Card("恢复与清理器", "执行回滚、恢复\n与环境清理", min_w=300, max_w=320).layout(measure),
    ]
    h_center(ctrl_cards, 360, 40, 240)
    ctrl_group = GroupBox("控制与调度层", COLORS["sand"], ctrl_cards, pad_t=110).layout(measure)

    exec_cards = [
        Card("宿主机注入器", "Kprobes / KVM / VFS / 内存管理链路", min_w=700, max_w=760).layout(measure),
        Card("Guest OS 注入器", "CPU 争抢 / 内存耗尽 / 网络异常 / 进程干预", min_w=700, max_w=760).layout(measure),
        Card("场景适配器", "k3s / Chaos Mesh 场景化注入", min_w=700, max_w=760).layout(measure),
    ]
    for idx, card in enumerate(exec_cards):
        card.x = 150
        card.y = 760 + idx * 220
    exec_group = GroupBox("故障注入执行层", COLORS["green"], exec_cards, pad_t=110).layout(measure)

    target_cards = [
        Card("宿主机与 KVM 环境", "ARM64 宿主机 / KVM / QEMU-KVM", min_w=640, max_w=700).layout(measure),
        Card("轻量级虚拟机集群", "Alpine 节点 / 控制节点 / 工作节点", min_w=640, max_w=700).layout(measure),
        Card("云原生上层业务", "k3s 工作负载 / Chaos Mesh 实验", min_w=640, max_w=700).layout(measure),
    ]
    for idx, card in enumerate(target_cards):
        card.x = 1150
        card.y = 760 + idx * 220
    target_group = GroupBox("目标系统与业务层", COLORS["purple"], target_cards, pad_t=110).layout(measure)

    user = Card("用户交互层", "Web 页面 / CLI / 参数配置", fill=COLORS["blue"], min_w=620, max_w=700).layout(measure)
    user.x = (ctrl_group.x1 + ctrl_group.x2 - user.w) // 2
    user.y = 90

    result = Card("结果与数据层", "日志回收、状态监控、实验结果存储与分析", fill=COLORS["blue"], min_w=1360, max_w=1420).layout(measure)
    result.x, result.y = 330, 1490

    for group in (ctrl_group, exec_group, target_group):
        draw_group(draw, group)
        add_bounds(bounds, group)
    for card in [user, result] + ctrl_cards + exec_cards + target_cards:
        draw_card(draw, card)
        add_bounds(bounds, card)

    ctrl_center_x = (ctrl_group.x1 + ctrl_group.x2) // 2
    exec_center_x = (exec_group.x1 + exec_group.x2) // 2
    target_center_x = (target_group.x1 + target_group.x2) // 2

    draw_arrow(draw, point_from(user, "bottom"), (ctrl_center_x, ctrl_group.y1))
    draw_arrow(draw, (exec_center_x, ctrl_group.y2), (exec_center_x, exec_group.y1))
    draw_arrow(draw, (target_center_x, ctrl_group.y2), (target_center_x, target_group.y1))

    for source, target_card in zip(exec_cards, target_cards):
        draw_arrow(draw, point_from(source, "right"), point_from(target_card, "left"))

    draw_arrow(draw, (exec_center_x, exec_group.y2), (exec_center_x, result.y))
    draw_arrow(draw, (target_center_x, target_group.y2), (target_center_x, result.y))
    crop_and_save(image, bounds, target)


def fig3_2(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    steps = [
        Card("参数配置", "选择故障类型、目标节点与持续时间", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure),
        Card("任务解析", "生成标准化任务描述与执行计划", fill=COLORS["sand"], min_w=760, max_w=820).layout(measure),
        Card("预检查", "验证 SSH 连通性、权限与环境状态", fill=COLORS["green"], min_w=760, max_w=820).layout(measure),
        Card("故障注入执行", "分发脚本并在目标节点触发注入", fill=COLORS["purple"], min_w=760, max_w=820).layout(measure),
        Card("状态监测与日志回收", "持续采样并收集系统行为变化", fill=COLORS["mint"], min_w=760, max_w=820).layout(measure),
        Card("恢复清理与结果输出", "回滚配置、结束实验并生成报告", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure),
    ]
    y = 100
    for step in steps:
        step.x = 700
        step.y = y
        y += step.h + 55
        draw_card(draw, step)
        add_bounds(bounds, step)
    for first, second in zip(steps, steps[1:]):
        draw_arrow(draw, point_from(first, "bottom"), point_from(second, "top"))
    crop_and_save(image, bounds, target)


def fig4_1(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    event = Card("客户机访问尚未映射的内存页", "触发硬件缺页异常 Page Fault", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure)
    event.x, event.y = 760, 80

    handler = Card("handle_mm_fault", "内核缺页处理函数", min_w=560, max_w=620).layout(measure)
    handler.x, handler.y = 890, 360

    inj = Card("注入分支", "Kretprobes 命中返回路径\n篡改故障码\n伪造 VM_FAULT_OOM", min_w=420, max_w=460).layout(measure)
    normal = Card("正常分支", "Stage-2 页表建立成功\n返回正常状态码", min_w=420, max_w=460).layout(measure)
    inj.x, inj.y = 520, 680
    normal.x, normal.y = 1180, 680

    bad = Card("异常结果", "客户机内存映射失败\n可能触发崩溃或内核恐慌", fill=COLORS["purple"], min_w=440, max_w=470).layout(measure)
    good = Card("正常结果", "控制权返回\n客户机继续运行", fill=COLORS["green"], min_w=440, max_w=470).layout(measure)
    bad.x, bad.y = 500, 1040
    good.x, good.y = 1200, 1040

    kernel_group = GroupBox("宿主机内核异常处理路径", COLORS["sand"], [handler, inj, normal, bad, good], pad_t=130, pad_b=70).layout(measure)
    draw_group(draw, kernel_group)
    add_bounds(bounds, kernel_group)

    for card in [event, handler, inj, normal, bad, good]:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(event, "bottom"), point_from(handler, "top"))
    split = (handler.x + handler.w // 2, handler.y + handler.h + 90)
    draw_poly_arrow(draw, [point_from(handler, "bottom"), split, point_from(inj, "top")])
    draw_poly_arrow(draw, [point_from(handler, "bottom"), split, point_from(normal, "top")])
    draw_arrow(draw, point_from(inj, "bottom"), point_from(bad, "top"))
    draw_arrow(draw, point_from(normal, "bottom"), point_from(good, "top"))
    crop_and_save(image, bounds, target)


def fig4_2(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    ctrl = Card("故障注入控制器", "根据实验参数选择目标进程与干预方式", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure)
    ctrl.x, ctrl.y = 700, 100

    signal = Card("POSIX 信号下发", "SIGSTOP / SIGCONT / SIGKILL", fill=COLORS["sand"], min_w=620, max_w=680).layout(measure)
    signal.x, signal.y = 770, 380

    states = [
        Card("暂停状态", "SIGSTOP 后进程挂起\n服务暂时失去响应", fill=COLORS["purple"], min_w=360, max_w=390).layout(measure),
        Card("恢复状态", "SIGCONT 后继续执行\n恢复原有控制流", fill=COLORS["green"], min_w=360, max_w=390).layout(measure),
        Card("终止状态", "SIGKILL 后立即退出\n触发上层恢复机制", fill=COLORS["mint"], min_w=360, max_w=390).layout(measure),
    ]
    h_center(states, 760, 70, 310)

    target_proc = Card("目标业务进程", "容器进程 / Pod / 应用服务", min_w=520, max_w=580).layout(measure)
    target_proc.x, target_proc.y = 820, 1100

    for card in [ctrl, signal, target_proc] + states:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(ctrl, "bottom"), point_from(signal, "top"))
    for state in states:
        draw_arrow(draw, point_from(signal, "bottom"), point_from(state, "top"))
        draw_arrow(draw, point_from(state, "bottom"), point_from(target_proc, "top"))
    crop_and_save(image, bounds, target)


def fig4_3(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    op = Card("控制节点", "通过 kubectl 下发 Chaos Mesh 实验", fill=COLORS["blue"], min_w=700, max_w=760).layout(measure)
    op.x, op.y = 730, 90

    host = Card("k3s 控制面", "接收 CRD 并完成实验对象编排", fill=COLORS["sand"], min_w=760, max_w=820).layout(measure)
    host.x, host.y = 700, 350

    systemvm = Card("Chaos Mesh 控制组件", "Controller Manager / Chaos Daemon", fill=COLORS["purple"], min_w=760, max_w=820).layout(measure)
    systemvm.x, systemvm.y = 700, 650

    observe = Card("目标 Pod 故障触发", "PodChaos / NetworkChaos / StressChaos 生效", fill=COLORS["green"], min_w=760, max_w=820).layout(measure)
    observe.x, observe.y = 700, 960

    recover = Card("恢复与验证", "删除实验资源并检查 Pod、事件与服务状态", fill=COLORS["mint"], min_w=760, max_w=820).layout(measure)
    recover.x, recover.y = 700, 1270

    for card in [op, host, systemvm, observe, recover]:
        draw_card(draw, card)
        add_bounds(bounds, card)
    for first, second in zip([op, host, systemvm, observe], [host, systemvm, observe, recover]):
        draw_arrow(draw, point_from(first, "bottom"), point_from(second, "top"))
    crop_and_save(image, bounds, target)


def fig4_4(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    steps = [
        Card("任务创建", "生成任务标识与故障描述", fill=COLORS["blue"], min_w=360, max_w=390).layout(measure),
        Card("环境预检查", "节点连通性 / 权限 / 初始状态", fill=COLORS["sand"], min_w=360, max_w=390).layout(measure),
        Card("故障触发", "脚本执行 / 内核探针 / 信号干预", fill=COLORS["green"], min_w=360, max_w=390).layout(measure),
        Card("运行监控", "状态采样 / 日志抓取 / 事件记录", fill=COLORS["purple"], min_w=360, max_w=390).layout(measure),
        Card("恢复清理", "结束干预、回滚配置并释放资源", fill=COLORS["mint"], min_w=360, max_w=390).layout(measure),
        Card("结果归档", "输出摘要、指标与实验记录", fill=COLORS["blue"], min_w=360, max_w=390).layout(measure),
    ]
    h_center(steps[:3], 220, 70, 220)

    run_monitor = steps[3]
    cleanup = steps[4]
    archive = steps[5]
    run_monitor.x = steps[2].x + (steps[2].w - run_monitor.w) // 2
    run_monitor.y = 760
    cleanup.x = run_monitor.x - cleanup.w - 70
    cleanup.y = 760
    archive.x = cleanup.x - archive.w - 70
    archive.y = 760

    for card in steps:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(steps[0], "right"), point_from(steps[1], "left"))
    draw_arrow(draw, point_from(steps[1], "right"), point_from(steps[2], "left"))
    draw_arrow(draw, point_from(steps[2], "bottom"), point_from(run_monitor, "top"))
    draw_arrow(draw, point_from(run_monitor, "left"), point_from(cleanup, "right"))
    draw_arrow(draw, point_from(cleanup, "left"), point_from(archive, "right"))
    crop_and_save(image, bounds, target)


def fig5_1(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    ui = Card("前端页面", "实验参数录入、运行状态展示、结果查看", fill=COLORS["blue"], min_w=760, max_w=820).layout(measure)
    ui.x, ui.y = 700, 90

    api = Card("FastAPI 接口层", "任务创建、状态查询、结果导出", fill=COLORS["sand"], min_w=760, max_w=820).layout(measure)
    api.x, api.y = 700, 340

    mid_cards = [
        Card("任务管理器", "解析配置\n维护任务生命周期", min_w=320, max_w=340).layout(measure),
        Card("调度执行器", "组织执行顺序\n控制远程下发", min_w=320, max_w=340).layout(measure),
        Card("日志与结果整理器", "采集日志\n生成实验摘要", min_w=360, max_w=380).layout(measure),
    ]
    h_center(mid_cards, 640, 50, 390)
    mid_group = GroupBox("控制中枢", COLORS["green"], mid_cards, pad_t=110).layout(measure)

    storage_cards = [
        Card("SSH 连接器", "与宿主机、虚拟机建立控制通道", fill=COLORS["paper"], min_w=480, max_w=520).layout(measure),
        Card("状态与结果存储", "缓存任务状态、日志摘要与报表", fill=COLORS["paper"], min_w=480, max_w=520).layout(measure),
    ]
    storage_cards[0].x, storage_cards[0].y = 420, 1120
    storage_cards[1].x, storage_cards[1].y = 1220, 1120
    storage_group = GroupBox("执行与数据支撑", COLORS["purple"], storage_cards, pad_t=110).layout(measure)

    for group in (mid_group, storage_group):
        draw_group(draw, group)
        add_bounds(bounds, group)
    for card in [ui, api] + mid_cards + storage_cards:
        draw_card(draw, card)
        add_bounds(bounds, card)

    draw_arrow(draw, point_from(ui, "bottom"), point_from(api, "top"))
    for card in mid_cards:
        draw_arrow(draw, point_from(api, "bottom"), point_from(card, "top"))
    draw_arrow(draw, point_from(mid_cards[0], "bottom"), point_from(storage_cards[0], "top"))
    draw_arrow(draw, point_from(mid_cards[2], "bottom"), point_from(storage_cards[1], "top"))
    crop_and_save(image, bounds, target)


def fig5_2(measure: Measure, target: Path) -> None:
    image, draw, bounds = canvas()
    ctrl_cards = [
        Card("前端页面", "故障参数配置与实验展示", min_w=380, max_w=410).layout(measure),
        Card("FastAPI 后端", "接口封装、任务生成与调度", min_w=380, max_w=410).layout(measure),
        Card("SSH 连接器", "远程下发命令与接收回执", min_w=380, max_w=410).layout(measure),
        Card("结果整理器", "生成实验日志与结果摘要", min_w=380, max_w=410).layout(measure),
    ]
    for idx, card in enumerate(ctrl_cards):
        card.x = 90
        card.y = 240 + idx * 220
    ctrl_group = GroupBox("控制节点", COLORS["sand"], ctrl_cards, pad_t=115).layout(measure)

    host_cards = [
        Card("宿主机 1", "Kprobes 模块 / KVM 路径 / VFS 注入", min_w=520, max_w=560).layout(measure),
        Card("宿主机 2", "备用宿主机 / 管理节点 / 监控对象", min_w=520, max_w=560).layout(measure),
    ]
    host_cards[0].x, host_cards[0].y = 980, 250
    host_cards[1].x, host_cards[1].y = 1570, 250
    host_group = GroupBox("宿主机侧执行域", COLORS["green"], host_cards, pad_t=115).layout(measure)

    vm_cards = [
        Card("VM 控制节点", "k3s server / API Server", min_w=360, max_w=390).layout(measure),
        Card("VM 工作节点 1", "k3s agent / 目标 Pod", min_w=360, max_w=390).layout(measure),
        Card("VM 工作节点 2", "k3s agent / 目标 Pod", min_w=360, max_w=390).layout(measure),
        Card("Chaos Mesh 组件", "Controller / Daemon / 实验 CRD 状态", min_w=640, max_w=700).layout(measure),
    ]
    vm_cards[0].x, vm_cards[0].y = 980, 860
    vm_cards[1].x, vm_cards[1].y = 1420, 860
    vm_cards[2].x, vm_cards[2].y = 1860, 860
    vm_cards[3].x, vm_cards[3].y = 1220, 1180
    vm_group = GroupBox("虚拟机侧执行域", COLORS["purple"], vm_cards, pad_t=115, pad_r=55).layout(measure)

    for group in (ctrl_group, host_group, vm_group):
        draw_group(draw, group)
        add_bounds(bounds, group)
    for card in ctrl_cards + host_cards + vm_cards:
        draw_card(draw, card)
        add_bounds(bounds, card)

    ssh = ctrl_cards[2]
    result = ctrl_cards[3]
    draw_poly_arrow(draw, [point_from(ssh, "right"), (760, ssh.y + ssh.h // 2), (760, host_cards[0].y + host_cards[0].h // 2), point_from(host_cards[0], "left")], width=4)
    draw_poly_arrow(draw, [point_from(ssh, "right"), (760, ssh.y + ssh.h // 2), (760, vm_cards[0].y + vm_cards[0].h // 2), point_from(vm_cards[0], "left")], width=4)
    draw_poly_arrow(draw, [point_from(result, "right"), (760, result.y + result.h // 2), (760, vm_cards[3].y + vm_cards[3].h // 2), point_from(vm_cards[3], "left")], width=4)
    crop_and_save(image, bounds, target)


FIGURES = {
    "图2-1": ("fig2_1.png", fig2_1),
    "图2-2": ("fig2_2.png", fig2_2),
    "图2-3": ("fig2_3.png", fig2_3),
    "图2-4": ("fig2_4.png", fig2_4),
    "图3-1": ("fig3_1.png", fig3_1),
    "图3-2": ("fig3_2.png", fig3_2),
    "图4-1": ("fig4_1.png", fig4_1),
    "图4-2": ("fig4_2.png", fig4_2),
    "图4-3": ("fig4_3.png", fig4_3),
    "图4-4": ("fig4_4.png", fig4_4),
    "图5-1": ("fig5_1.png", fig5_1),
    "图5-2": ("fig5_2.png", fig5_2),
}


def replace_images(doc_path: Path, mapping: dict[str, Path]) -> None:
    if not doc_path.exists():
        return

    backup = doc_path.with_name(doc_path.stem + "-修图前备份.docx")
    if not backup.exists():
        shutil.copy2(doc_path, backup)

    doc = Document(doc_path)
    paragraphs = doc.paragraphs

    for idx, para in enumerate(paragraphs):
        caption = para.text.strip()
        for prefix, image_path in mapping.items():
            if caption.startswith(prefix):
                picture_para = paragraphs[idx - 1]
                elem = picture_para._p
                for child in list(elem):
                    elem.remove(child)
                run = picture_para.add_run()
                run.add_picture(str(image_path), width=Cm(14.2))
                picture_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                break

    doc.save(doc_path)


def build_preview(image_paths: list[Path]) -> None:
    thumbs = []
    for path in image_paths:
        im = Image.open(path).convert("RGB")
        im.thumbnail((540, 380))
        canvas_im = Image.new("RGB", (560, 410), "white")
        canvas_im.paste(im, ((560 - im.width) // 2, 20))
        d = ImageDraw.Draw(canvas_im)
        font = ImageFont.truetype(FONT_MEDIUM, 24)
        d.text((20, 370), path.stem, font=font, fill=COLORS["text"])
        thumbs.append(canvas_im)

    cols = 3
    rows = math.ceil(len(thumbs) / cols)
    grid = Image.new("RGB", (cols * 560, rows * 410), "#F8FAFC")
    for idx, thumb in enumerate(thumbs):
        x = (idx % cols) * 560
        y = (idx // cols) * 410
        grid.paste(thumb, (x, y))
    grid.save(PREVIEW_PATH, format="PNG")


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    measure = Measure()
    generated: dict[str, Path] = {}

    for prefix, (filename, builder) in FIGURES.items():
        target = OUT_DIR / filename
        builder(measure, target)
        generated[prefix] = target

    build_preview(list(generated.values()))
    for doc_path in DOCS:
        replace_images(doc_path, generated)

    print("generated", len(generated))
    for key, path in generated.items():
        print(key, path)
    print("preview", PREVIEW_PATH)


if __name__ == "__main__":
    main()
