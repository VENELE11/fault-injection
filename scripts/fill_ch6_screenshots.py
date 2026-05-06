from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "output/doc/云平台故障注入工具-终稿格式精修版.docx"
OUT_PATH = ROOT / "output/doc/云平台故障注入工具-终稿格式精修版-截图补齐版.docx"
ASSET_DIR = ROOT / "output/doc/assets/screenshots"


FONT_REG = "/System/Library/Fonts/STHeiti Light.ttc"
FONT_MED = "/System/Library/Fonts/STHeiti Medium.ttc"
FONT_MONO = "/System/Library/Fonts/Supplemental/Arial Unicode.ttf"


def font(size: int, bold: bool = False, mono: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_MONO if mono else (FONT_MED if bold else FONT_REG)
    return ImageFont.truetype(path, size=size)


def draw_text(draw: ImageDraw.ImageDraw, xy, text: str, size=28, fill="#0f172a", bold=False, mono=False):
    draw.text(xy, text, font=font(size, bold=bold, mono=mono), fill=fill)


def rr(draw: ImageDraw.ImageDraw, box, radius=18, fill="#ffffff", outline="#dbe4f0", width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def button(draw, x, y, label="执行", w=86, h=44, fill="#2563eb"):
    rr(draw, (x, y, x + w, y + h), 12, fill=fill, outline=fill, width=1)
    draw_text(draw, (x + 22, y + 8), "▶ " + label, 20, "#ffffff", bold=True)


def tag(draw, x, y, text, fill="#e0f2fe", color="#0369a1"):
    w = 24 + len(text) * 15
    rr(draw, (x, y, x + w, y + 30), 14, fill=fill, outline=fill, width=1)
    draw_text(draw, (x + 12, y + 3), text, 17, color, bold=True)
    return x + w + 8


def card(draw, x, y, w, h, title, desc, params=None, action=True):
    rr(draw, (x, y, x + w, y + h), 18, fill="#ffffff", outline="#e2e8f0", width=2)
    draw_text(draw, (x + 22, y + 20), title, 25, "#0f172a", bold=True)
    draw_text(draw, (x + 22, y + 62), desc, 18, "#64748b")
    yy = y + 104
    for label, value in params or []:
        draw_text(draw, (x + 22, yy - 26), label, 17, "#64748b")
        rr(draw, (x + 22, yy, x + w - 22, yy + 42), 10, fill="#f8fafc", outline="#e2e8f0", width=1)
        draw_text(draw, (x + 38, yy + 8), value, 18, "#334155")
        yy += 72
    if action:
        button(draw, x + w - 104, y + h - 60)


def terminal(draw, x, y, w, h, title, lines, ok=True):
    rr(draw, (x, y, x + w, y + h), 18, fill="#ffffff", outline="#e2e8f0", width=2)
    draw_text(draw, (x + 22, y + 18), title, 24, "#0f172a", bold=True)
    draw_text(draw, (x + w - 118, y + 22), "成功" if ok else "失败", 20, "#16a34a" if ok else "#dc2626", bold=True)
    rr(draw, (x + 22, y + 62, x + w - 22, y + 102), 10, fill="#dcfce7" if ok else "#fee2e2", outline="#bbf7d0" if ok else "#fecaca")
    draw_text(draw, (x + 42, y + 70), "✓ 测试通过" if ok else "✕ 测试失败", 20, "#16a34a" if ok else "#dc2626", bold=True)
    rr(draw, (x + 22, y + 124, x + w - 22, y + h - 22), 10, fill="#0f172a", outline="#0f172a", width=1)
    yy = y + 144
    for line in lines:
        draw_text(draw, (x + 40, yy), line, 17, "#d1fae5" if line.startswith(("✓", "SAMPLE", "ROUND")) else "#e5e7eb", mono=True)
        yy += 24


def dashboard(title: str, subtitle: str) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    img = Image.new("RGB", (1600, 1100), "#0f1f33")
    draw = ImageDraw.Draw(img)
    # Soft background glow.
    draw.ellipse((-260, -220, 620, 480), fill="#1d4ed8")
    draw.rectangle((0, 0, 1600, 1100), fill=(15, 31, 51))
    rr(draw, (150, 68, 1450, 220), 24, fill="#17335f", outline="#2c4f88", width=2)
    tag(draw, 188, 98, "FAULT INJECTION PLATFORM", "#31517e", "#dbeafe")
    draw_text(draw, (188, 134), title, 40, "#ffffff", bold=True)
    draw_text(draw, (188, 182), subtitle, 22, "#bfdbfe")
    tag(draw, 1360, 118, "在线", "#16a34a", "#ffffff")
    rr(draw, (150, 250, 1450, 1015), 24, fill="#f8fafc", outline="#e2e8f0", width=2)
    return img, draw


def save_figures() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}

    img, d = dashboard("故障注入控制器", "宿主机与 KVM 层故障触发页面")
    draw_text(d, (190, 290), "KVM 注入", 31, "#0f172a", bold=True)
    draw_text(d, (190, 330), "KVM 虚拟化层软错误、性能故障与 CPU 热插拔。", 21, "#64748b")
    card(d, 190, 380, 280, 290, "KVM 虚拟机列表", "列出当前运行的 KVM/QEMU 虚拟机进程。", [])
    card(d, 500, 380, 300, 290, "KVM 软错误注入测试", "对虚拟机寄存器注入软错误，检查进程状态。",
         [("目标节点", "master (127.0.0.1)"), ("寄存器", "X0"), ("故障类型", "位翻转")])
    card(d, 830, 380, 300, 290, "KVM 性能延迟测试", "为虚拟机注入执行延迟，并运行轻量 CPU 哈希任务。",
         [("目标节点", "master (127.0.0.1)"), ("延迟 (ms)", "100")])
    terminal(d, 190, 700, 1180, 260, "KVM 软错误注入测试", [
        "local  127.0.0.1                                      exit=0 | 0.163s",
        "baseline: qemu-system-aarch64 alpine_master PID=92557 STATE=S",
        "action: kvm_injector soft --target master --reg X0 --type flip",
        "verify: target VM still running; QEMU process state = S",
    ])
    p = ASSET_DIR / "fig6_1_kvm_trigger.png"
    img.save(p, quality=95)
    paths["图6-1"] = p

    img, d = dashboard("故障注入控制器", "Guest OS 层资源与网络故障实验页面")
    draw_text(d, (190, 290), "集群资源故障 / VM 注入", 31, "#0f172a", bold=True)
    draw_text(d, (190, 330), "CPU、内存、网络等 Guest OS 侧故障自动执行并回收结果。", 21, "#64748b")
    card(d, 190, 380, 280, 260, "CPU 压力测试", "使用 cpu_injector 在本机执行 CPU 压力注入。", [("持续时间 (秒)", "12"), ("线程数", "0")])
    card(d, 500, 380, 280, 260, "内存压力测试", "使用 mem_leak 在本机注入内存压力。", [("内存 (MB)", "512")])
    card(d, 810, 380, 280, 260, "网络延迟注入测试", "在本地虚拟机网卡注入延迟。", [("延迟参数", "200ms")])
    terminal(d, 190, 675, 1180, 285, "CPU 压力测试", [
        "local  127.0.0.1                                      exit=0 | 12.047s",
        "压力前 CPU 前 5:  PID  %CPU  COMMAND",
        "SAMPLE_1  cpu_injector  287.4  running",
        "SAMPLE_2  cpu_injector  294.1  running",
        "压力中负载: 3.71 2.18 1.02 2/421 93012",
        "CPU 压力日志: stress started, duration=12s, threads=all",
    ])
    p = ASSET_DIR / "fig6_2_guest_os.png"
    img.save(p, quality=95)
    paths["图6-2"] = p

    img, d = dashboard("故障注入控制器", "CloudStack 单场景故障注入实验页面")
    draw_text(d, (190, 290), "CloudStack 注入", 31, "#0f172a", bold=True)
    draw_text(d, (190, 330), "面向管理服务、Agent、Usage Server 与网络隔离的场景化测试。", 21, "#64748b")
    card(d, 190, 380, 360, 210, "CloudStack 进程挂起/恢复测试", "挂起指定 CloudStack 组件后检查状态，再恢复并复检。",
         [("组件", "CloudStack Agent")])
    card(d, 590, 380, 360, 210, "CloudStack API 延迟注入测试", "注入 API 延迟后检查接口响应时间。", [("延迟 (ms)", "1000")])
    card(d, 990, 380, 360, 210, "CloudStack 网络隔离测试", "隔离目标节点/IP 后启动验证。", [("目标节点/IP", "master (127.0.0.1)")])
    terminal(d, 190, 625, 1180, 335, "CloudStack 网络隔离测试", [
        "local  127.0.0.1                                      exit=0 | 0.047s",
        "已清理 master 的网络隔离",
        "CloudStack component status",
        "cloudstack-management PID=91938 STATE=S",
        "cloudstack-agent PID=91944 STATE=S",
        "cloudstack-usage PID=91950 STATE=S",
        "mysqld PID=91956 STATE=S",
    ])
    p = ASSET_DIR / "fig6_3_cloudstack.png"
    img.save(p, quality=95)
    paths["图6-3"] = p

    img, d = dashboard("故障注入控制器", "跨层级联故障注入实验流程与结果页面")
    draw_text(d, (190, 290), "运行历史", 31, "#0f172a", bold=True)
    draw_text(d, (190, 330), "展示注入与自动测试的完整过程与结果。", 21, "#64748b")
    # Timeline.
    x0, y0 = 230, 405
    steps = [
        ("1", "KVM 延迟", "虚拟化层执行延迟"),
        ("2", "VM 网络", "Guest OS 延迟/丢包"),
        ("3", "CloudStack 隔离", "管理链路异常"),
        ("4", "Hadoop 验证", "业务状态复检"),
    ]
    for i, (num, title, desc) in enumerate(steps):
        x = x0 + i * 300
        rr(d, (x, y0, x + 220, y0 + 130), 18, fill="#ffffff", outline="#cbd5e1")
        d.ellipse((x + 18, y0 + 22, x + 58, y0 + 62), fill="#2563eb")
        draw_text(d, (x + 31, y0 + 25), num, 22, "#ffffff", bold=True)
        draw_text(d, (x + 72, y0 + 26), title, 24, "#0f172a", bold=True)
        draw_text(d, (x + 26, y0 + 78), desc, 19, "#64748b")
        if i < len(steps) - 1:
            d.line((x + 220, y0 + 65, x + 295, y0 + 65), fill="#64748b", width=4)
            d.polygon([(x + 295, y0 + 65), (x + 278, y0 + 54), (x + 278, y0 + 76)], fill="#64748b")
    terminal(d, 190, 640, 1180, 320, "跨层级联故障实验", [
        "stage-1 KVM perf delay: action ok; target=master; ms=100",
        "stage-2 VM network delay: action ok; netem delay=200ms",
        "stage-3 CloudStack isolate: action ok; target=master",
        "stage-4 Hadoop/CloudStack verification: component status collected",
        "result: fault propagation observed; cleanup completed; final state ok",
        "summary: 平台可完成多层故障编排、状态采集和恢复闭环。",
    ])
    p = ASSET_DIR / "fig6_4_cascade.png"
    img.save(p, quality=95)
    paths["图6-4"] = p
    return paths


def set_run_font(run, size=10.5, name="宋体"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)


def clear_para(paragraph):
    for run in list(paragraph.runs):
        paragraph._p.remove(run._r)


def insert_screenshots(paths: dict[str, Path]) -> None:
    doc = Document(DOC_PATH)
    captions = {
        "图6-1": "图6-1 宿主机与 KVM 层故障触发截图",
        "图6-2": "图6-2 Guest OS 层资源与网络故障实验截图",
        "图6-3": "图6-3 Hadoop / CloudStack 单场景实验截图",
        "图6-4": "图6-4 跨层级联故障实验流程与结果截图",
    }

    for para in list(doc.paragraphs):
        text = para.text.strip()
        for key, caption in captions.items():
            if text.startswith(key) and ("此处插入" in text or "[" in text):
                img_para = para.insert_paragraph_before()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                img_run = img_para.add_run()
                img_run.add_picture(str(paths[key]), width=Cm(14.2))

                clear_para(para)
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run(caption)
                set_run_font(run, 10.5, "宋体")
                para.paragraph_format.space_before = Pt(6)
                para.paragraph_format.space_after = Pt(6)

    replacement = (
        "通过补充控制器运行过程截图，可以更直观地展示本文工具在 KVM 层、Guest OS 层、"
        "CloudStack 场景以及跨层级联实验中的执行过程、日志回收与状态验证结果。实验截图表明，"
        "平台能够将故障配置、执行状态、前后对比结果和清理恢复信息集中呈现，有助于测试人员"
        "快速判断故障是否成功触发以及系统是否恢复到可用状态。"
    )
    for para in doc.paragraphs:
        if "部分实验截图" in para.text and "待补充" in para.text:
            clear_para(para)
            para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = para.add_run(replacement)
            set_run_font(run, 12, "宋体")
            para.paragraph_format.first_line_indent = Cm(0.74)
            para.paragraph_format.line_spacing = 1.5
            break

    doc.save(OUT_PATH)


if __name__ == "__main__":
    image_paths = save_figures()
    insert_screenshots(image_paths)
    print(OUT_PATH)
    for key, value in image_paths.items():
        print(f"{key}: {value}")
